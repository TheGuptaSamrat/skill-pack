#!/usr/bin/env python3
"""Extract and analyze content from FPSL .docx document"""

import sys
import json
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import parse_xml
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'], check=True)
    from docx import Document

def extract_document(docx_path):
    """Extract full document content"""
    doc = Document(docx_path)
    
    output = {
        'metadata': {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'sections': len(doc.sections),
        },
        'sections': []
    }
    
    current_section = None
    
    # Extract all paragraphs with style info
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        style = para.style.name if para.style else 'Normal'
        level = para.paragraph_format.outline_level
        
        item = {
            'type': 'paragraph',
            'text': text,
            'style': style,
            'level': level,
        }
        
        # Track section headers
        if style.startswith('Heading'):
            current_section = text
            item['is_section'] = True
        
        output['sections'].append(item)
    
    # Extract all tables
    for table_num, table in enumerate(doc.tables, 1):
        table_data = {
            'type': 'table',
            'number': table_num,
            'rows': len(table.rows),
            'columns': len(table.columns),
            'content': []
        }
        
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data['content'].append(row_data)
        
        output['sections'].append(table_data)
    
    return output

if __name__ == '__main__':
    docx_file = '/Users/samratgupta/skill-pack/metadata-drop/fpsl/Building SAP FPSL_FSDM Golden Source.docx'
    
    print("Extracting document...")
    data = extract_document(docx_file)
    
    print(f"\n📄 Document Statistics:")
    print(f"   Paragraphs: {data['metadata']['paragraphs']}")
    print(f"   Tables: {data['metadata']['tables']}")
    print(f"   Sections: {data['metadata']['sections']}")
    
    # Print structure
    print(f"\n📋 Document Structure:\n")
    headers = [s for s in data['sections'] if s.get('is_section')]
    for h in headers[:20]:
        print(f"  {'  ' * h['level']}→ {h['text'][:80]}")
    
    # Print full content to markdown
    md_path = '/Users/samratgupta/skill-pack/metadata-drop/fpsl/extracted-content.md'
    with open(md_path, 'w') as f:
        for item in data['sections']:
            if item['type'] == 'paragraph':
                text = item['text']
                level = item['level']
                style = item['style']
                
                if style.startswith('Heading'):
                    f.write(f"{'#' * (level + 1)} {text}\n\n")
                elif style.startswith('List'):
                    f.write(f"- {text}\n")
                else:
                    f.write(f"{text}\n\n")
            elif item['type'] == 'table':
                f.write(f"\n### Table {item['number']}\n\n")
                for row in item['content']:
                    f.write(f"| {' | '.join(row)} |\n")
                f.write("\n")
    
    print(f"\n✅ Extracted content saved to: {md_path}")
    print(f"   Total items extracted: {len(data['sections'])}")

